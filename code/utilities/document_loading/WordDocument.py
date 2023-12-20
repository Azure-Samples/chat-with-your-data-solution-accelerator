from typing import List
from io import BytesIO
from docx import Document
import requests
from .DocumentLoadingBase import DocumentLoadingBase
from ..common.SourceDocument import SourceDocument


class WordDocumentLoading(DocumentLoadingBase):
    """
    This class represents a Word document loader that converts the content of a Word document into Markdown format.

    Attributes:
        doc_headings_to_markdown_tags (dict): A dictionary mapping Word document heading styles to Markdown tags.
    """

    def __init__(self) -> None:
        super().__init__()
        self.doc_headings_to_markdown_tags = {
            'Heading 1': 'h1',
            'Heading 2': 'h2',
            'Heading 3': 'h3',
            'Heading 4': 'h4',
            'Heading 5': 'h5',
            'Heading 6': 'h6',
        }

    def _download_document(self, document_url: str) -> BytesIO:
        """
        Download the Word document from the given URL.

        Args:
            document_url (str): The URL of the Word document.

        Returns:
            BytesIO: The downloaded Word document as a BytesIO object.
        """
        response = requests.get(document_url)
        file = BytesIO(response.content)
        return file

    def _get_opening_tag(self, heading_level: int) -> str:
        """
        Get the opening markdown tag for the specified heading level.

        Args:
            heading_level (int): The level of the heading.

        Returns:
            str: The opening markdown tag.
        """
        return f"<{self.doc_headings_to_markdown_tags.get(f'{heading_level}', '')}>"

    def _get_closing_tag(self, heading_level: int) -> str:
        """
        Get the closing markdown tag for the specified heading level.

        Args:
            heading_level (int): The level of the heading.

        Returns:
            str: The closing markdown tag.
        """
        return f"</{self.doc_headings_to_markdown_tags.get(f'{heading_level}', '')}>"

    def load(self, document_url: str) -> List[SourceDocument]:
        """
        Load the Word document from the given URL and convert it into markdown format.

        Args:
            document_url (str): The URL of the Word document.

        Returns:
            List[SourceDocument]: A list of SourceDocument objects containing the converted markdown content.
        """
        output = ""
        document = Document(self._download_document(document_url))
        for paragraph in document.paragraphs:
            output += f"{self._get_opening_tag(paragraph.style.name)}{paragraph.text}{self._get_closing_tag(paragraph.style.name)}\n"
        documents = [
            SourceDocument(
                content=output,
                source=document_url,
                offset=0,
                page_number=0,
            )
        ]
        return documents
