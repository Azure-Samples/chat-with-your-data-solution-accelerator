from typing import List
from io import BytesIO
from docx import Document
import requests
from .document_loading_base import DocumentLoadingBase
from ..common.source_document import SourceDocument


class WordDocumentLoading(DocumentLoadingBase):
    def __init__(self) -> None:
        super().__init__()
        self.doc_headings_to_markdown_tags = {
            "Heading 1": "h1",
            "Heading 2": "h2",
            "Heading 3": "h3",
            "Heading 4": "h4",
            "Heading 5": "h5",
            "Heading 6": "h6",
        }

    def _download_document(self, document_url: str) -> BytesIO:
        response = requests.get(document_url)
        file = BytesIO(response.content)
        return file

    def _get_opening_tag(self, heading_level: int) -> str:
        return f"<{self.doc_headings_to_markdown_tags.get(f'{heading_level}', 'p')}>"

    def _get_closing_tag(self, heading_level: int) -> str:
        return f"</{self.doc_headings_to_markdown_tags.get(f'{heading_level}', 'p')}>"

    def load(self, document_url: str) -> List[SourceDocument]:
        output = ""
        document = Document(self._download_document(document_url))
        for paragraph in document.paragraphs:
            output += f"{self._get_opening_tag(paragraph.style.name)}{paragraph.text}{self._get_closing_tag(paragraph.style.name)}\n"
        documents = [
            SourceDocument(
                content=output,
                source=document_url,
                offset=0,
                page_number=0,
            )
        ]
        return documents
